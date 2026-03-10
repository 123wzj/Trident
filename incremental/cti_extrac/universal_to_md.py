import os
import re
from pathlib import Path
from urllib.parse import urlparse
from typing import Union

import requests
import html2text
from docx import Document

from agent.env_utils import MINERU_API_KEY


class UniversalToMarkdown:
    """统一将多种格式转换为干净 Markdown，适配大模型输入"""

    def __init__(self, use_mineru_for_web: bool = True, default_save_dir: str = "./cti"):
        """
        :param use_mineru_for_web: 是否优先使用 MinerU 处理网页（推荐 True）
        :param default_save_dir: 默认保存目录
        """
        self.use_mineru_for_web = use_mineru_for_web
        self.default_save_dir = default_save_dir

    def convert(self, source: Union[str, Path], save_to_dir: str = None) -> str:
        """
        统一入口：自动识别输入类型并转换为 Markdown
        :param source: 可以是 URL、本地文件路径（str 或 Path）
        :param save_to_dir: 可选,保存目录(如 './cti')，None 则使用默认目录
        :return: 清洗后的 Markdown 字符串
        """
        source = str(source)

        # 判断输入类型并转换
        if self._is_url(source):
            md_result = self._from_url(source)
            # 传入转换后的内容以提取标题
            source_name = self._url_to_filename(source, content=md_result)
        elif os.path.isfile(source):
            md_result = self._from_file(source)
            source_name = Path(source).stem
        else:
            raise ValueError(f"Invalid input: {source} is neither a valid URL nor a file path")

        # 自动保存(使用指定目录或默认目录)
        target_dir = save_to_dir if save_to_dir is not None else self.default_save_dir
        if target_dir:
            self._save_markdown(md_result, source_name, target_dir)

        return md_result

    def _url_to_filename(self, url: str, content: str = None) -> str:
        """从 Markdown 内容提取标题作为文件名"""
        import hashlib

        # 从 Markdown 提取标题
        if content:
            lines = content.strip().split('\n')
            for line in lines[:20]:  # 检查前20行
                line = line.strip()
                if line.startswith('# '):
                    title = line.lstrip('#').strip()
                    return self._sanitize_filename(title)
                # 跳过空行和其他 Markdown 语法
                if line and not line.startswith(('*', '-', '>', '```', '|', '#', '[', '!')):
                    # 取第一个实质性段落的前80字符作为标题
                    title = line[:80]
                    return self._sanitize_filename(title)

        # 降级：用 URL 最后一部分
        parsed = urlparse(url)
        path_parts = [p for p in parsed.path.split('/') if p]
        if path_parts:
            return self._sanitize_filename(path_parts[-1])

        # 最后降级：域名 + hash
        url_hash = hashlib.md5(url.encode()).hexdigest()[:8]
        return f"{parsed.netloc}_{url_hash}"

    def _sanitize_filename(self, title: str, max_length: int = 100) -> str:
        """清理为合法文件名"""
        # 移除非法字符和多余空格
        title = re.sub(r'[<>:"/\\|?*\n\r]', '', title)
        title = re.sub(r'\s+', '_', title.strip())
        title = re.sub(r'[^\w\-_.]', '', title)

        # 限制长度并移除末尾的特殊字符
        if len(title) > max_length:
            title = title[:max_length].rstrip('_-.')

        return title if title else "untitled"

    def _save_markdown(self, content: str, filename: str, save_dir: str):
        """保存 Markdown 到指定目录"""
        save_path = Path(save_dir)
        save_path.mkdir(parents=True, exist_ok=True)

        output_file = save_path / f"{filename}.md"
        output_file.write_text(content, encoding='utf-8')

        print(f"已保存到: {output_file.absolute()}")

    def _is_url(self, s: str) -> bool:
        try:
            result = urlparse(s)
            return all([result.scheme, result.netloc])
        except Exception:
            return False

    def _from_url(self, url: str) -> str:
        """处理网页 URL（优先 MinerU API，备用 html2text）"""
        if self.use_mineru_for_web and MINERU_API_KEY and not url.endswith('/'):
            try:
                print("使用 MinerU API 提取...")
                md = self._mineru_api_convert(url, "markdown")
                print(f"MinerU 提取成功: {len(md)} 字符")
                return self._post_clean(md)
            except Exception as e:
                print(f"MinerU API 失败: {e}")
                print("   切换到 html2text 提取...")

        # 备用方案：html2text (完整提取)
        try:
            md = self._html2text_extract(url)
            print(f"html2text 提取成功: {len(md)} 字符")
            return self._post_clean(md)
        except Exception as e:
            raise RuntimeError(f"Failed to process URL {url}: {e}")

    def _html2text_extract(self, url: str) -> str:
        """使用 html2text 提取完整网页内容"""
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36'
        }

        response = requests.get(url, headers=headers, timeout=30)
        response.raise_for_status()

        # 配置 html2text
        h = html2text.HTML2Text()
        h.ignore_links = False          # 保留链接
        h.ignore_images = True         # 不保留图片
        h.ignore_emphasis = False       # 保留强调（粗体/斜体）
        h.ignore_tables = False         # 保留表格
        h.body_width = 0                # 不自动换行
        h.skip_internal_links = False   # 保留内部链接
        h.inline_links = True           # 使用内联链接格式
        h.wrap_links = False            # 不包装链接
        h.default_image_alt = "image"   # 图片默认 alt 文本

        md = h.handle(response.text)
        return md

    def _from_file(self, filepath: str) -> str:
        """处理本地文件"""
        ext = os.path.splitext(filepath)[1].lower()

        if ext == ".pdf":
            return self._from_pdf(filepath)
        elif ext == ".docx":
            return self._from_docx(filepath)
        elif ext == ".html":
            return self._from_html(filepath)
        elif ext == ".txt":
            return self._from_txt(filepath)
        else:
            raise ValueError(f"Unsupported file format: {ext}")

    def _from_pdf(self, filepath: str) -> str:
        """PDF 转 Markdown（使用 MinerU 官方 API）"""
        if not MINERU_API_KEY:
            raise RuntimeError("MinerU API key not configured. Cannot process PDF.")

        # 1. 如果是本地文件，先上传到 transfer.sh 获取临时 URL
        if not self._is_url(filepath):
            print(f"上传 PDF 到临时存储...")
            url = self._upload_to_transfer_sh(filepath)
        else:
            url = filepath  # 已是 URL

        # 2. 调用 MinerU 官方 API
        print(f"使用 MinerU API 处理 PDF...")
        md = self._mineru_api_convert(url, "markdown")
        print(f"PDF 转换成功: {len(md)} 字符")
        return self._post_clean(md)

    def _mineru_api_convert(self, source: str, output_format: str = "markdown") -> str:
        """调用 MinerU 官方 API 转换内容"""
        API_URL = "https://api.mineru.net/v1/extract"

        payload = {
            "url": source,
            "output_format": output_format,
            "model_version": "vlm"  # 推荐使用 vlm 模型（含版面分析）
        }

        headers = {
            "Authorization": f"Bearer {MINERU_API_KEY}",
            "Content-Type": "application/json"
        }

        response = requests.post(API_URL, json=payload, headers=headers, timeout=60)

        if response.status_code == 200:
            result = response.json()
            if result.get("status") == "success":
                return result["data"]["content"]
            else:
                raise RuntimeError(f"MinerU API error: {result.get('message')}")
        else:
            raise RuntimeError(f"API request failed (HTTP {response.status_code}): {response.text}")

    def _upload_to_transfer_sh(self, filepath: str) -> str:
        """将本地文件上传到 transfer.sh 获取临时公开 URL"""
        with open(filepath, 'rb') as f:
            response = requests.post('https://transfer.sh/', files={'file': f}, timeout=60)
            if response.status_code == 200:
                url = response.text.strip()
                print(f"上传成功: {url}")
                return url
            else:
                raise RuntimeError(f"Failed to upload to transfer.sh: {response.text}")

    def _from_docx(self, filepath: str) -> str:
        """Word 转 Markdown"""
        try:
            doc = Document(filepath)
            full_text = []

            for para in doc.paragraphs:
                if para.text.strip():
                    # 简单转换标题（根据样式）
                    if para.style.name.startswith('Heading'):
                        level = min(int(para.style.name[-1]), 6) if para.style.name[-1].isdigit() else 2
                        full_text.append(f"{'#' * level} {para.text}")
                    else:
                        full_text.append(para.text)

            # 表格处理
            for table in doc.tables:
                md_table = self._table_to_md(table)
                if md_table:
                    full_text.append("\n" + md_table + "\n")

            md = "\n\n".join(full_text)
            return self._post_clean(md)
        except Exception as e:
            raise RuntimeError(f"Failed to parse DOCX {filepath}: {e}")

    def _from_html(self, filepath: str) -> str:
        """HTML 文件转 Markdown（使用 html2text）"""
        with open(filepath, 'r', encoding='utf-8') as f:
            html_content = f.read()

        # 配置 html2text
        h = html2text.HTML2Text()
        h.ignore_links = False
        h.ignore_images = False
        h.ignore_emphasis = False
        h.ignore_tables = False
        h.body_width = 0
        h.inline_links = True

        md = h.handle(html_content)
        return self._post_clean(md)

    def _from_txt(self, filepath: str) -> str:
        """纯文本转 Markdown（仅基础清洗）"""
        with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
            text = f.read()
        # 按空行分段
        paragraphs = [p.strip() for p in text.split('\n\n') if p.strip()]
        return '\n\n'.join(paragraphs)

    def _table_to_md(self, table) -> str:
        """将 python-docx 表格转为 Markdown 表格"""
        try:
            rows = []
            for i, row in enumerate(table.rows):
                cells = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
                rows.append('| ' + ' | '.join(cells) + ' |')
                if i == 0:
                    separator = '| ' + ' | '.join(['---'] * len(cells)) + ' |'
                    rows.append(separator)
            return '\n'.join(rows)
        except Exception:
            return ""

    def _post_clean(self, md: str, highlight_ioc: bool = False) -> str:
        """后处理：进一步清洗噪声，适配大模型
        :param highlight_ioc: 是否高亮 IOC(默认关闭,避免破坏格式)
        """
        if not md:
            return ""

        # 1. 移除多余空行（保留最多2个连续空行）
        md = re.sub(r'\n{3,}', '\n\n', md)

        # 2. 移除可能残留的调试信息
        md = re.sub(r'<!--.*?-->', '', md, flags=re.DOTALL)

        # 3. 移除页眉页脚特征词（如 "Page 1 of 10"）
        md = re.sub(r'Page \d+ of \d+', '', md, flags=re.IGNORECASE)

        # 4. IOC 高亮(可选,且仅在非代码块区域)
        if highlight_ioc:
            md = self._highlight_ioc_safe(md)

        return md.strip()

    def _highlight_ioc_safe(self, md: str) -> str:
        """安全地高亮 IOC,避免破坏现有格式"""
        lines = []
        in_code_block = False

        for line in md.split('\n'):
            # 跳过代码块
            if line.strip().startswith('```'):
                in_code_block = not in_code_block
                lines.append(line)
                continue

            if not in_code_block and '`' not in line:
                # 仅在非代码行高亮
                # IP: 更严格的匹配(避免版本号)
                line = re.sub(
                    r'(?<!\d)(\d{1,3}\.\d{1,3}\.\d{1,3}\.\d{1,3})(?!\d)',
                    lambda m: f'`IP:{m.group(1)}`' if self._is_valid_ip(m.group(1)) else m.group(1),
                    line
                )
                # SHA256
                line = re.sub(r'\b([a-fA-F0-9]{64})\b', r'`HASH:\1`', line)

            lines.append(line)

        return '\n'.join(lines)

    def _is_valid_ip(self, ip: str) -> bool:
        """验证是否为有效 IP 地址"""
        parts = ip.split('.')
        try:
            return all(0 <= int(part) <= 255 for part in parts)
        except ValueError:
            return False


# ======================
# 使用示例
# ======================
if __name__ == "__main__":
    # 初始化转换器，默认保存到 ./cti 目录
    converter = UniversalToMarkdown(use_mineru_for_web=True, default_save_dir="./cti")

    # 测试网页 URL（自动保存到 ./cti）
    # url = "https://www.crowdstrike.com/en-us/blog/fake-recovery-manual-used-to-deliver-unidentified-stealer/"
    # url = "https://thedfirreport.com/2025/06/30/hide-your-rdp-password-spray-leads-to-ransomhub-deployment/"
    url = "https://www.welivesecurity.com/en/eset-research/evasive-panda-leverages-monlam-festival-target-tibetans/"
    md_output = converter.convert(url)


    # 测试本地文件
    # txt_path = "./test_samples/threat_report.txt"
    # if os.path.exists(txt_path):
    #     print(f"\n\n 测试本地文件: {txt_path}")
    #     md_output = converter.convert(txt_path)
    #     print(f"TXT 转换成功！")

    # docx_path = "./test_samples/threat_report.docx"
    # if os.path.exists(docx_path):
    #     print(f"\n\n测试 DOCX 文件: {docx_path}")
    #     md_output = converter.convert(docx_path)
    #     print(f"DOCX 转换成功！")